"""Word / PDF 上传解析为 Tiptap HTML，并写入 SQLite 文档。"""

from __future__ import annotations

import html
import io
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile
from PyPDF2 import PdfReader
from sqlalchemy.orm import Session

from docx.oxml.ns import qn

from db.models import Document as DbDocument
from utils.file_utils import sanitize_text
from utils.page_break_html import PAGE_BREAK_HTML


def _escape(text: str) -> str:
    return html.escape(sanitize_text(text), quote=True)


def filename_to_title(filename: str) -> str:
    stem = Path(filename or "导入文档").stem.strip()
    title = sanitize_text(stem) or "导入文档"
    return title[:512]


def detect_upload_kind(filename: str, raw: bytes) -> str:
    """
    根据扩展名与文件头判断类型。
    返回 'docx' | 'pdf'，无法识别时抛出 HTTPException。
    """
    ext = Path(filename or "").suffix.lower()

    if ext == ".doc":
        raise HTTPException(
            status_code=422,
            detail="不支持旧版 .doc 格式，请在 Word 中「另存为」.docx 后再导入",
        )
    if ext in (".docx", ".pdf"):
        return ext.lstrip(".")

    header = raw[:8]
    if raw.startswith(b"%PDF"):
        return "pdf"
    if raw.startswith(b"PK"):
        return "docx"

    raise HTTPException(
        status_code=422,
        detail=(
            f"无法识别文件类型（{filename or '无文件名'}）。"
            "请上传 .docx 或 .pdf；勿使用旧版 .doc"
        ),
    )


def _paragraph_style_name(paragraph) -> str:
    try:
        return paragraph.style.name if paragraph.style else ""
    except Exception:
        return ""


def _paragraph_has_manual_page_break(paragraph) -> bool:
    """Word 手动分页符（run 内 w:br type=page）或段前分页。"""
    try:
        if paragraph.paragraph_format.page_break_before:
            return True
    except Exception:
        pass
    for run in paragraph.runs:
        for br in run._element.findall(qn("w:br")):
            br_type = br.get(qn("w:type"))
            if br_type is None or br_type == "page":
                return True
    return False


def _is_list_paragraph(paragraph) -> bool:
    name = _paragraph_style_name(paragraph)
    if "List" in name:
        return True
    try:
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.numPr is not None:
            return True
    except Exception:
        pass
    return False


def docx_bytes_to_html(data: bytes) -> str:
    """从 .docx 提取段落并转为 Tiptap 友好 HTML。"""
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"无法读取 Word 文件（请确认是有效的 .docx）：{e}",
        ) from e

    blocks: list[str] = []
    list_buffer: list[str] = []

    def flush_list() -> None:
        nonlocal list_buffer
        if not list_buffer:
            return
        items = "".join(f"<li><p>{_escape(line)}</p></li>" for line in list_buffer)
        blocks.append(f"<ul>{items}</ul>")
        list_buffer = []

    for paragraph in doc.paragraphs:
        has_break = _paragraph_has_manual_page_break(paragraph)
        text = sanitize_text(paragraph.text)

        if text:
            style = _paragraph_style_name(paragraph)
            if _is_list_paragraph(paragraph):
                list_buffer.append(text)
            else:
                flush_list()
                if "Heading 1" in style or style.startswith("Title"):
                    blocks.append(f"<h1>{_escape(text)}</h1>")
                elif "Heading 2" in style:
                    blocks.append(f"<h2>{_escape(text)}</h2>")
                elif "Heading 3" in style:
                    blocks.append(f"<h3>{_escape(text)}</h3>")
                else:
                    blocks.append(f"<p>{_escape(text)}</p>")
        elif not has_break:
            flush_list()
            continue

        if has_break:
            flush_list()
            blocks.append(PAGE_BREAK_HTML)

    flush_list()

    if not blocks:
        return "<p>（未识别到段落文本，请在编辑器中补充内容）</p>"
    return "".join(blocks)


def _pdf_page_to_paragraphs(page_text: str) -> list[str]:
    raw = sanitize_text(page_text)
    if not raw:
        return []
    chunks = [c.strip() for c in re.split(r"\n\s*\n+", raw) if c.strip()]
    if len(chunks) > 1:
        return chunks
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    return lines if lines else [raw]


def pdf_bytes_to_html(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"无法读取 PDF 文件：{e}",
        ) from e

    blocks: list[str] = []
    total_pages = len(reader.pages)
    for index, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        paragraphs = _pdf_page_to_paragraphs(page_text)
        if paragraphs:
            for para in paragraphs:
                blocks.append(f"<p>{_escape(para)}</p>")
        elif total_pages > 1:
            blocks.append(
                f"<p><em>（第 {index + 1} 页未提取到文本，可能是扫描页）</em></p>"
            )
        if total_pages > 1 and index < total_pages - 1:
            blocks.append(PAGE_BREAK_HTML)

    if not blocks:
        return "<p>（未能从 PDF 提取文本，可能是扫描件；请在编辑器中补充内容）</p>"
    return "".join(blocks)


def bytes_to_tiptap_html(data: bytes, kind: str) -> str:
    if kind == "docx":
        return docx_bytes_to_html(data)
    if kind == "pdf":
        return pdf_bytes_to_html(data)
    raise HTTPException(status_code=422, detail="不支持的文件类型")


async def create_document_from_upload(
    file: UploadFile,
    db: Session,
) -> DbDocument:
    filename = (file.filename or "").strip()
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=422, detail="空文件，请选择有效文档")

    kind = detect_upload_kind(filename, raw)
    content_html = bytes_to_tiptap_html(raw, kind)
    title = filename_to_title(filename) if filename else f"导入_{kind}"
    now = datetime.now(timezone.utc)

    doc = DbDocument(
        title=title,
        content=content_html,
        created_at=now,
        updated_at=now,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc
